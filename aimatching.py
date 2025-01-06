import datetime
import streamlit as st
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
import json
from google.oauth2 import service_account
from googleapiclient.discovery import build
import openai
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.errors import HttpError
import gdown
import tempfile
import os
import google.auth.transport.requests
import concurrent.futures
from io import BytesIO

# OpenAI APIキーをStreamlitのシークレットから設定
openai.api_key = st.secrets["openai_api_key"]

# サービスアカウント認証情報
SERVICE_ACCOUNT_FILE = 'service_account.json'
SCOPES = [
    'https://www.googleapis.com/auth/spreadsheets',
    'https://www.googleapis.com/auth/documents',
    'https://www.googleapis.com/auth/drive'
]

# Google API認証
def authenticate_google_services():
    with open(SERVICE_ACCOUNT_FILE, 'r') as f:
        service_account_info = json.load(f)
    credentials = service_account.Credentials.from_service_account_info(
        service_account_info, scopes=SCOPES
    )
    return credentials

# Google Docs APIでドキュメントの内容を取得する関数
def get_document_text(docs_service, document_id):
    try:
        document = docs_service.documents().get(documentId=document_id).execute()
        content = document.get('body').get('content')
        document_text = ''
        for element in content:
            if 'paragraph' in element:
                for text_run in element.get('paragraph').get('elements'):
                    if 'textRun' in text_run:
                        document_text += text_run.get('textRun').get('content')
        return document_text
    except Exception as e:
        st.error(f"ドキュメントの読み取り中にエラーが発生しました: {e}")
        return ""

# 一週間以内で、D列とE列にフリーワードが含まれている件を取得する関数
def get_filtered_resent_sheet_data(sheet_service, sheet_id, range_name, search_query):
    try:
        sheet = sheet_service.spreadsheets()
        result = sheet.values().get(spreadsheetId=sheet_id, range=range_name).execute()
        values = result.get('values', [])

        #テスト用  
        # st.write("values (first 2 rows):", values[:2])
        
        # 今日の日付と1週間前の日付を計算
        today = datetime.date.today()
        one_week_ago = today - datetime.timedelta(days=7)
        
        filtered_data = []
        
        for row in values:
            try:
                data_datetime = datetime.datetime.strptime(row[0], "%Y/%m/%d %H:%M:%S")
                data_date = data_datetime.date()

                # 一週間内のデータだけフィルタリング
                if one_week_ago <= data_date <= today:
                
                    # D列とE列にフリーワードが含まれているデータだけフィルタリング
                    if search_query.lower() in row[4].lower():
                        filtered_data.append(row)
            except Exception as e:
            
                # 日付やデータ形式にエラーがある場合はスキップ
                continue  
                
    except Exception as e:
        st.error(f"スプレッドシートデータ取得中にエラーが発生しました: {e}")
        return []
        
    return filtered_data

# フリーワード検索結果の中でF列がTRUEまたはFALSEで直近50件を取得する関数
def get_sheet_filtered_data(filtered_data, filter_value):
    try:
        
        # 指定された列の条件で行をフィルタリング
        filtered_values = [row for row in filtered_data if len(row) > 5 and row[5].strip().lower() == filter_value.lower()]
        recent_values = filtered_values[:50]
        
        #テスト用  
        # st.write("recent_values (first 2 rows):", recent_values[:2])
        
        return recent_values
    except Exception as e:
        st.error(f"スプレッドシートデータ取得中にエラーが発生しました: {e}")
        return []
        
# F列がTRUEまたはFALSEで直近50件を取得する関数
def get_sheet_data(sheet_service, sheet_id, range_name, filter_column, filter_value):
    try:
        sheet = sheet_service.spreadsheets()
        result = sheet.values().get(spreadsheetId=sheet_id, range=range_name).execute()
        values = result.get('values', [])
        
        # 指定された列の条件で行をフィルタリング
        filtered_values = [row for row in values if len(row) > filter_column and row[filter_column].strip().lower() == filter_value]
        recent_values = filtered_values[:50]
        return recent_values
    except Exception as e:
        st.error(f"スプレッドシートデータ取得中にエラーが発生しました: {e}")
        return []

# Google Drive APIでファイル内容を取得する関数
def get_file_content_from_drive(file_id):
    extracted_text = []  # 추출된 텍스트를 저장할 리스트 초기화

    link = f"https://drive.google.com/uc?id={file_id}"

    # 파일 다운로드 시도
    try:
        file_path = download_file_from_drive(link)
        if not file_path:
            extracted_text.append("")  # 다운로드 실패 시 빈 문자열 추가
            st.warning("ファイルのダウンロードに失敗しました。")
            return extracted_text
    except Exception as e:
        st.error(f"ダウンロードエラー: {e}")
        extracted_text.append("")  # 오류 발생 시 빈 문자열 추가
        return extracted_text

    # 파일 다운로드 성공 후 처리
    st.success(f"ファイルをダウンロードしました: {file_path}")
    if file_path.endswith(".pdf"):
        st.info("PDF ファイルとして解析中...")
        pdf_text = extract_text_from_pdf(file_path)
        extracted_text.append(pdf_text if pdf_text else "")  # 추출된 텍스트 저장
        if pdf_text:
            st.text_area("抽出されたテキスト", pdf_text, height=300)
    elif file_path.endswith((".xls", ".xlsx")):
        st.info("Excel ファイルとして解析中...")
        excel_text = extract_text_from_excel(file_path)
        extracted_text.append(excel_text if excel_text else "")  # 추출된 텍스트 저장
        if excel_text:
            st.text_area("抽出されたテキスト", excel_text, height=300)
    else:
        st.warning("対応していないファイル形式です。PDF または Excel ファイルを使用してください。")
        extracted_text.append("")  # 지원하지 않는 파일 형식일 경우 빈 문자열 추가

    return extracted_text  # 추출된 텍스트 반환


# テキスト抽出関数
def extract_text_from_file(uploaded_file):
    file_type = uploaded_file.type
    extracted_text = ""

    if file_type == "application/pdf":
        pdf_reader = PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            extracted_text += page.extract_text()
    elif file_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
        df = pd.read_excel(uploaded_file)
        extracted_text = df.to_string()
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        for paragraph in doc.paragraphs:
            extracted_text += paragraph.text + "\n"
    else:
        st.error("対応しているファイル形式はPDF、Excel、Wordです。")
    return extracted_text

# NG企業シートのA列データ取得関数。
def get_ng_companies(sheet_service, spreadsheet_id):
    try:
        # NG企業のA列データを取得
        ng_result = sheet_service.spreadsheets().values().get(
            spreadsheetId=spreadsheet_id, range="'NG企業'!A2:A"
        ).execute()
        ng_values = ng_result.get('values', [])
        
        # データをフラットなリストに変換し、小文字化
        ng_list = [
            row[0].strip().lower() for row in ng_values if row
            and row[0].strip().lower() != "フリーランス" #NG企業からフリーランスは追加しない
        ]
        return ng_list
    except Exception as e:
        st.error(f"NG企業データの取得中にエラーが発生しました: {e}")
        return []

# フリーワード検索関数
def search_data(sheet_id, range_name, form_id):
        
    # Google APIサービスの認証
    credentials = authenticate_google_services()
    sheet_service = build('sheets', 'v4', credentials=credentials)

    form_key = f'search_form_{form_id}'
    
    if f"search_query_{form_id}" not in st.session_state:
        st.session_state[f"search_query_{form_id}"] = ""
    if form_id not in st.session_state:
        st.session_state[form_id] = []
    if f"search_message_{form_id}" not in st.session_state:
        st.session_state[f"search_message_{form_id}"] = ""

    if range_name == "'【案件】DB'!A:O":
        filter_value = "false"
    elif range_name == "'【人材】DB'!A:O":
        filter_value = "true"
    else:
        filter_value = None
        
    with st.form(key=form_key):
        st.text("フリーワード検索による絞り込み")
        col1, col2 = st.columns([3, 1]) 
        with col1:
            search_query = st.text_input("", placeholder="フリーワード", label_visibility="collapsed", value=st.session_state[f"search_query_{form_id}"])
        with col2:
            submit_button = st.form_submit_button("チェック")  
            
        if submit_button:
            try:
                st.session_state[f"search_query_{form_id}"] = search_query
                
                sheet = sheet_service.spreadsheets()
                result = sheet.values().get(spreadsheetId=sheet_id, range=range_name).execute()
                values = result.get('values', [])

                # 今日の日付と1週間前の日付を計算
                today = datetime.date.today()
                one_week_ago = today - datetime.timedelta(days=7)
                
                filtered_values = []
                
                filter_column = 5  
                
                for row in values[1:]:
                    try:
                        # 9列目の値が「TRUE」の場合、NG企業として除外
                        is_ng_company = len(row) > 9 and row[9].strip().lower() == "true"
                        
                        if (
                            len(row) > filter_column and 
                            row[filter_column].strip().lower() == filter_value and
                            not is_ng_company and
                            one_week_ago <= datetime.datetime.strptime(row[0], "%Y/%m/%d").date() <= today and
                            search_query.lower() in row[4].lower()
                        ):
                            filtered_values.append(row)
                    except (ValueError, IndexError) as e:
                        st.write(f"Exception caught: {e}")
                
                st.session_state[form_id] = filtered_values
                st.session_state[f"search_message_{form_id}"] = f"「{st.session_state[f'search_query_{form_id}']}」を含む案件が {len(st.session_state[form_id])} 件見つかりました！"
                
            except Exception as e:
                st.error(f"Google Sheets データ検索中にエラーが発生しました: {e}")

        if st.session_state[f"search_message_{form_id}"]:
            st.success(st.session_state[f"search_message_{form_id}"])

def process_and_render_results(raw_response):
    parsed_results = raw_response.split("\n---\n")

    if parsed_results:
        st.markdown(parsed_results[0])  
    
    for idx, result in enumerate(parsed_results[1:], start=1):  
        st.markdown(result) 
        
        
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button(f"コピー {idx}", key=f"copy_{idx}"):
                st.session_state["clipboard"] = result
                st.success(f"{idx}件目の内容をコピーしました！")
        with col2:
            if st.button(f"メール {idx}", key=f"mail_{idx}"):
                st.session_state["mail_template"] = result
                st.success(f"{idx}件目のメールを準備中！")

def extract_file_id(file_url):
    """
    Google Drive URL からファイル ID を抽出する関数。
    """
    try:
        match = re.search(r"/d/([a-zA-Z0-9_-]+)", file_url)
        if match:
            return match.group(1)
        else:
            raise ValueError(f"URLからファイルIDを抽出できませんでした: {file_url}")
    except Exception as e:
        raise ValueError(f"ファイルID抽出中にエラーが発生しました: {e}")

def download_file_from_drive(link):
    try:
        file_path = gdown.download(link, quiet=False)
        if file_path:
            file_name = os.path.basename(file_path)
            return file_name
    except Exception as e:
        st.error(f"ファイルのダウンロードに失敗しました: {e}")
        return None

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        st.error(f"PDF のテキスト抽出に失敗しました: {e}")
        return None

def extract_text_from_excel(file_path):
    extracted_text = []  # extracted_text 리스트를 함수 시작 부분에서 초기화
    try:
        # 파일 확장자에 따라 처리 엔진을 선택
        if file_path.endswith(".xls"):
            engine = "xlrd"  # .xls 파일의 경우 xlrd 사용
        elif file_path.endswith(".xlsx"):
            engine = "openpyxl"  # .xlsx 파일의 경우 openpyxl 사용
        else:
            raise ValueError("지원하지 않는 파일 형식입니다. .xls 또는 .xlsx 파일을 사용하세요.")

        # 모든 시트를 데이터프레임으로 읽음
        df_list = pd.read_excel(file_path, sheet_name=None)
        for sheet_name, df in df_list.items():
            try:
                # 비어 있는 행과 열 제거
                df_cleaned = df.dropna(how='all').dropna(axis=1, how='all')
                
                # NaN 값을 공백("")으로 대체
                df_cleaned = df_cleaned.fillna("")
                
                # 데이터가 존재하는 경우만 추가
                if not df_cleaned.empty:
                    extracted_text.append(f"【シート名: {sheet_name}】\n")
                    extracted_text.append(df_cleaned.to_string(index=False, header=False))  # 열 이름 출력 안 함
                    extracted_text.append("\n\n")
                else:
                    extracted_text.append("")  # 데이터가 없으면 빈 문자열 추가
            except Exception as e:
                st.warning(f"シート {sheet_name} の解析中にエラーが発生しました: {e}")
                extracted_text.append("")  # 해당 시트 처리 실패 시 빈 문자열 추가
    except Exception as e:
        st.error(f"Excel ファイル全体の解析中にエラーが発生しました: {e}")
        extracted_text.append("")  # 전체 오류 시 빈 문자열 추가
    
    # 결과 텍스트를 합쳐 반환
    return "\n".join(extracted_text)



# def main():
#     extracted_text = []  # 추출된 텍스트를 저장할 리스트 초기화

#     st.title("Google ドライブファイルテキスト抽出ツール")
#     st.write("Google ドライブの共有リンクを入力してください。")
#     link = st.text_input("共有リンク", "")

#     if st.form_submit_button("ダウンロードとテキスト抽出"):
#         if not link:
#             st.warning("リンクを入力してください。")
#             extracted_text.append("")  # 링크가 없으면 빈 문자열 추가
#             return extracted_text
#         else:
#             st.info("ファイルをダウンロード中...")

#             # 파일 다운로드 시도
#             try:
#                 file_path = download_file_from_drive(link)
#                 if not file_path:
#                     extracted_text.append("")  # 다운로드 실패 시 빈 문자열 추가
#                     st.warning("ファイルのダウンロードに失敗しました。")
#                     return extracted_text
#             except Exception as e:
#                 st.error(f"ダウンロードエラー: {e}")
#                 extracted_text.append("")  # 오류 발생 시 빈 문자열 추가
#                 return extracted_text

#             # 파일 다운로드 성공 후 처리
#             st.success(f"ファイルをダウンロードしました: {file_path}")
#             if file_path.endswith(".pdf"):
#                 st.info("PDF ファイルとして解析中...")
#                 pdf_text = extract_text_from_pdf(file_path)
#                 extracted_text.append(pdf_text if pdf_text else "")  # 추출된 텍스트 저장
#                 if pdf_text:
#                     st.text_area("抽出されたテキスト", pdf_text, height=300)
#             elif file_path.endswith((".xls", ".xlsx")):
#                 st.info("Excel ファイルとして解析中...")
#                 excel_text = extract_text_from_excel(file_path)
#                 extracted_text.append(excel_text if excel_text else "")  # 추출된 텍스트 저장
#                 if excel_text:
#                     st.text_area("抽出されたテキスト", excel_text, height=300)
#             else:
#                 st.warning("対応していないファイル形式です。PDF または Excel ファイルを使用してください。")
#                 extracted_text.append("")  # 지원하지 않는 파일 형식일 경우 빈 문자열 추가

#     return extracted_text  # 추출된 텍스트 반환

def process_file(file_id, drive_service):
    """
    Google Drive 파일을 처리하여 텍스트를 추출하는 함수.

    Args:
        file_id (str): Google Drive 파일 ID.
        drive_service: Google Drive API 서비스 객체.
    
    Returns:
        dict: 파일 이름과 추출된 텍스트를 포함하는 결과.
    """
    try:
        # 파일 메타데이터 가져오기
        file_metadata = drive_service.files().get(
            fileId=file_id,
            fields="mimeType, name",
            supportsAllDrives=True
        ).execute()
        mime_type = file_metadata["mimeType"]
        file_name = file_metadata["name"]
        st.write(f"File Metadata: {file_metadata}")

        # 파일 처리
        if mime_type == "application/pdf":
            # PDF 파일 처리
            request = drive_service.files().get_media(fileId=file_id)
            file_content = BytesIO(request.execute())
            pdf_reader = PdfReader(file_content)
            extracted_text = "".join([page.extract_text() for page in pdf_reader.pages])
            st.write(f"Extracted Text from {file_name} (Preview):")
            st.text(extracted_text[:500])

        elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
            # Excel 파일 처리
            request = drive_service.files().get_media(fileId=file_id)
            file_content = BytesIO(request.execute())
            df = pd.read_excel(file_content)
            st.write(f"Extracted DataFrame from {file_name}:")
            st.write(df)
            extracted_text = df.to_string(index=False)
            st.write(f"Extracted Text from {file_name} (Preview):")
            st.text(extracted_text[:500])

        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Word 파일 처리
            request = drive_service.files().get_media(fileId=file_id)
            file_content = BytesIO(request.execute())
            doc = Document(file_content)
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
            st.write(f"Extracted Text from {file_name} (Preview):")
            st.text(extracted_text[:500])

        else:
            # 지원하지 않는 형식
            raise ValueError(f"지원하지 않는 MIME 타입: {mime_type}")

        return {
            "file_name": file_name,
            "extracted_text": extracted_text
        }

    except Exception as e:
        st.warning(f"파일 처리 중 오류가 발생했습니다. 파일 ID: {file_id}, 에러: {e}")
        return {
            "file_name": None,
            "extracted_text": None
        }

def process_multiple_files(file_ids, drive_service):
    """
    여러 Google Drive 파일 ID를 처리하여 텍스트를 추출하는 함수.

    Args:
        file_ids (str): 쉼표로 구분된 Google Drive 파일 ID 문자열.
        drive_service: Google Drive API 서비스 객체.

    Returns:
        list: 각 파일 이름과 추출된 텍스트를 포함하는 딕셔너리의 리스트.
    """
    # 쉼표로 구분된 파일 ID를 리스트로 변환
    file_id_list = [file_id.strip() for file_id in file_ids.split(",") if file_id.strip()]
    results = []

    for file_id in file_id_list:
        try:
            # 파일 메타데이터 가져오기
            file_metadata = drive_service.files().get(
                fileId=file_id,
                fields="mimeType, name",
                supportsAllDrives=True
            ).execute()
            mime_type = file_metadata["mimeType"]
            file_name = file_metadata["name"]
            st.write(f"File Metadata for {file_id}: {file_metadata}")

            # 파일 처리
            if mime_type == "application/pdf":
                # PDF 파일 처리
                request = drive_service.files().get_media(fileId=file_id)
                file_content = BytesIO(request.execute())
                pdf_reader = PdfReader(file_content)
                extracted_text = "".join([page.extract_text() for page in pdf_reader.pages])
                st.write(f"Extracted Text from {file_name} (Preview):")
                st.text(extracted_text[:500])

            elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
                # Excel 파일 처리
                request = drive_service.files().get_media(fileId=file_id)
                file_content = BytesIO(request.execute())
                df = pd.read_excel(file_content)
                st.write(f"Extracted DataFrame from {file_name}:")
                st.write(df)
                extracted_text = df.to_string(index=False)
                st.write(f"Extracted Text from {file_name} (Preview):")
                st.text(extracted_text[:500])

            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Word 파일 처리
                request = drive_service.files().get_media(fileId=file_id)
                file_content = BytesIO(request.execute())
                doc = Document(file_content)
                extracted_text = "\n".join([para.text for para in doc.paragraphs])
                st.write(f"Extracted Text from {file_name} (Preview):")
                st.text(extracted_text[:500])

            else:
                # 지원하지 않는 형식
                raise ValueError(f"지원하지 않는 MIME 타입: {mime_type}")

            results.append({
                "file_name": file_name,
                "extracted_text": extracted_text
            })

        except Exception as e:
            st.warning(f"파일 처리 중 오류가 발생했습니다. 파일 ID: {file_id}, 에러: {e}")
            results.append({
                "file_name": f"Error processing {file_id}",
                "extracted_text": None
            })

    return results

        
def extract_text_from_drive_file(drive_service, file_id):
    """
    Google Drive API를 사용하여 파일에서 텍스트를 추출합니다.
    PDF와 Excel 파일 형식에 따라 자동 처리.
    """
    try:
        # 파일 메타데이터 가져오기 (MIME 타입 확인)
        file_metadata = drive_service.files().get(fileId=file_id, fields="*").execute()
        st.write(file_metadata)
        mime_type = file_metadata["mimeType"]
        file_name = file_metadata["name"]
        
        st.write(f"Processing file: {file_name} (MIME type: {mime_type})")
        
        # PDF 파일 처리
        if mime_type == "application/pdf":
            request = drive_service.files().get_media(fileId=file_id)
            file_content = BytesIO(request.execute())
            
            # PyPDF2를 사용하여 텍스트 추출
            from PyPDF2 import PdfReader
            pdf_reader = PdfReader(file_content)
            extracted_text = ""
            for page in pdf_reader.pages:
                extracted_text += page.extract_text()

            # 첫 500자만 출력
            st.write("Extracted Text (Preview):", extracted_text[:500])
            return extracted_text

        # Excel 파일 처리
        elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
            request = drive_service.files().get_media(fileId=file_id)
            file_content = BytesIO(request.execute())
            
            # Pandas로 Excel 데이터 읽기
            import pandas as pd
            df = pd.read_excel(file_content)
            extracted_text = df.to_string(index=False)
            
            # 첫 500자만 출력
            st.write("Extracted Text (Preview):", extracted_text[:500])
            return extracted_text

        # 지원하지 않는 형식
        else:
            raise ValueError(f"지원하지 않는 MIME 타입: {mime_type}")

    except Exception as e:
        raise RuntimeError(f"파일 처리 중 오류 발생: {e}")

# 1. 다운로드 없이 Google Drive에서 직접 텍스트 추출
def extract_text_from_drive(file_id, drive_service):
    try:
        file_metadata = drive_service.files().get(
                fileId=file_id, 
                fields="mimeType, name", 
                supportsAllDrives=True  # Shared Drive 지원
            ).execute()
        # Debug: 파일 메타데이터 확인
        st.write(f"Processing File ID: {file_id}, Name: {file_name}, MimeType: {mime_type}")
    except Exception as e:
        st.error(f"Error fetching metadata for File ID {file_id}: {str(e)}")
        return f"Error fetching metadata for File ID {file_id}: {str(e)}"

        mime_type = file_metadata['mimeType']
        file_name = file_metadata['name']
        st.write(f"File Name: {file_name}, MimeType: {mime_type}")

        try:
            request = drive_service.files().get_media(fileId=file_id)
            file_stream = io.BytesIO()
            downloader = MediaIoBaseDownload(file_stream, request)
                
            done = False
            while not done:
                _, done = downloader.next_chunk()
                
            file_stream.seek(0)
        except Exception as e:
            st.error(f"Error downloading file content: {str(e)}")
            return f"Error downloading file content: {str(e)}"
            
        # PDF 처리
        if mime_type == "application/pdf":
            reader = PdfReader(file_stream)
            pdf_text = "\n".join(page.extract_text() for page in reader.pages)
            st.write("PDF file processed successfully.")
            return pdf_text

        # Excel 처리
        elif mime_type in [
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ]:
            try:
                excel_data = pd.read_excel(file_stream, sheet_name=None)  # 모든 시트 읽기
                extracted_text = []
                for sheet_name, sheet_data in excel_data.items():
                    extracted_text.append(f"Sheet: {sheet_name}\n{sheet_data.to_string(index=False, header=True)}")
                st.write("Excel file processed successfully.")
                return "\n\n".join(extracted_text)
            except Exception as e:
                st.error(f"Error processing Excel file: {str(e)}")
                return f"Error processing Excel file: {str(e)}"


        # 지원하지 않는 파일 형식 처리
        else:
            return f"지원하지 않는 파일 형식입니다: {mime_type}"

    except Exception as e:
        st.write(f"Error processing file ID {file_id}: {e}")  # Debug 에러 로그
        return f"Error processing file ID {file_id}: {e}"

    # results = {}

    # # 각 파일 처리
    # for file_id in file_ids:
    #     results[file_id] = process_file(file_id)

    # return results
    

def extract_text_from_drive_file(drive_service, file_id):
    """
    Google Drive 파일에서 텍스트를 추출하는 함수.
    PDF와 Excel 파일 형식에 따라 자동 처리.
    """
    try:
        # 파일 메타데이터 가져오기 (MIME 타입 확인)
        file_metadata = drive_service.files().get(fileId=file_id, fields="mimeType, name").execute()
        mime_type = file_metadata["mimeType"]
        file_name = file_metadata["name"]
        
        print(f"Processing file: {file_name} (MIME type: {mime_type})")
        
        # PDF 파일 처리
        if mime_type == "application/pdf":
            request = drive_service.files().get_media(fileId=file_id)
            file_content = BytesIO(request.execute())
            
            # PDF 텍스트 추출
            pdf_reader = PdfReader(file_content)
            extracted_text = ""
            for page in pdf_reader.pages:
                extracted_text += page.extract_text()
            return extracted_text

        # Excel 파일 처리
        elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
            request = drive_service.files().get_media(fileId=file_id)
            file_content = BytesIO(request.execute())
            
            # Pandas로 Excel 데이터 읽기
            df = pd.read_excel(file_content)
            extracted_text = df.to_string(index=False)
            return extracted_text

        # 지원하지 않는 형식
        else:
            raise ValueError(f"지원하지 않는 MIME 타입: {mime_type}")

    except Exception as e:
        raise RuntimeError(f"파일 처리 중 오류 발생: {e}")
        
def extract_file_text_by_ids(file_ids):
    """
    Google Drive 파일 ID가 쉼표로 구분된 경우, 각각의 파일 내용을 추출하고 결합합니다.
    """
    extracted_texts = []  # 여러 파일의 텍스트를 저장할 리스트

    # 쉼표로 분리된 ID 처리
    file_id_list = [file_id.strip() for file_id in file_ids.split(",") if file_id.strip()]

    for file_id in file_id_list:
        try:
            # 개별 ID에 대해 텍스트 추출
            link = f"https://drive.google.com/uc?id={file_id}"
            file_path = download_file_from_drive(link)

            if not file_path:
                raise ValueError(f"ファイルのダウンロードに失敗しました: {file_id}")

            # 파일 타입에 따라 텍스트 추출
            if file_path.endswith(".pdf"):
                extracted_text = extract_text_from_pdf(file_path)
            elif file_path.endswith((".xls", ".xlsx")):
                extracted_text = extract_text_from_excel(file_path)
            else:
                raise ValueError(f"対応していないファイル形式: {file_id}")
            
            extracted_texts.append(extracted_text)
        except Exception as e:
            extracted_texts.append(f"エラーが発生しました ({file_id}): {e}")
        finally:
            # 파일 삭제
            if file_path and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception as e:
                    st.warning(f"一時ファイルの削除中にエラーが発生しました ({file_id}): {e}")

    # 모든 파일의 내용을 결합
    return "\n\n".join(extracted_texts)



# def main2(file_id):
#     extracted_text = []  # 추출된 텍스트를 저장할 리스트 초기화

#     link = f"https://drive.google.com/uc?id={file_id}"

#     # 파일 다운로드 시도
#     try:
#         file_path = download_file_from_drive(link)
#         if not file_path:
#             extracted_text.append("")  # 다운로드 실패 시 빈 문자열 추가
#             st.warning("ファイルのダウンロードに失敗しました。")
#             return extracted_text
#     except Exception as e:
#         st.error(f"ダウンロードエラー: {e}")
#         extracted_text.append("")  # 오류 발생 시 빈 문자열 추가
#         return extracted_text

#     # 파일 다운로드 성공 후 처리
#     st.success(f"ファイルをダウンロードしました: {file_path}")
#     if file_path.endswith(".pdf"):
#         st.info("PDF ファイルとして解析中...")
#         pdf_text = extract_text_from_pdf(file_path)
#         extracted_text.append(pdf_text if pdf_text else "")  # 추출된 텍스트 저장
#         if pdf_text:
#             st.text_area("抽出されたテキスト", pdf_text, height=300)
#     elif file_path.endswith((".xls", ".xlsx")):
#         st.info("Excel ファイルとして解析中...")
#         excel_text = extract_text_from_excel(file_path)
#         extracted_text.append(excel_text if excel_text else "")  # 추출된 텍스트 저장
#         if excel_text:
#             st.text_area("抽出されたテキスト", excel_text, height=300)
#     else:
#         st.warning("対応していないファイル形式です。PDF または Excel ファイルを使用してください。")
#         extracted_text.append("")  # 지원하지 않는 파일 형식일 경우 빈 문자열 추가

#     return extracted_text  # 추출된 텍스트 반환

def validate_service_account_key(key_file_path):
    try:
        credentials = service_account.Credentials.from_service_account_file(key_file_path)
        request = google.auth.transport.requests.Request()
        credentials.refresh(request)
        st.write("서비스 계정 키가 유효합니다.")
    except Exception as e:
        st.write(f"서비스 계정 키가 유효하지 않습니다: {e}")

# メイン処理
if "password_correct" not in st.session_state:
    st.session_state.password_correct = False

if "selected_tool" not in st.session_state:
    st.session_state.selected_tool = None

# ユーザーネーム
username = "AIMATCHING001"

# 認証処理
if not st.session_state.password_correct:
    st.title("ログイン")
    st.markdown(f"### ユーザー名: {username}")
    password = st.text_input("パスワードを入力してください", type="password")
    if password == "ai":
        st.session_state.password_correct = True
        st.success("認証に成功しました！")
        st.rerun()
    elif password:
        st.error("パスワードが間違っています。")

# 認証後の処理
else:
    # Streamlitアプリケーション
    st.title("AI MATCHING")
    
    # タブの作成
    tab1, tab2, tab3, tab4 = st.tabs(["人材要件から最適案件詳細検索","案件簡易検索", "案件概要から最適人材詳細検索", "人材簡易検索"])
    
    # NG企業リストを初期化 
    # if "ng_list" in st.session_state:
    credentials = authenticate_google_services()
    sheet_service = build('sheets', 'v4', credentials=credentials)
    #     st.session_state["ng_list"] = get_ng_companies(sheet_service, "1amJJDVMr3__OmLgWo1Z9w6FXZ9aMaNm0WRlx1_TYXnE")
    
    with tab1:
        search_data(
            sheet_id="1amJJDVMr3__OmLgWo1Z9w6FXZ9aMaNm0WRlx1_TYXnE",
            range_name="'【案件】DB'!A:O",
            form_id="tab1"
        )
        
        # 情報入力フォーム
        with st.form("input_form_detail"):
            # main()
            # validate_service_account_key(SERVICE_ACCOUNT_FILE)
            st.subheader("情報入力フォーム")
                            
            # エンジニアスキル
            skills = st.text_input("エンジニアスキル (カンマ区切り)", placeholder="例: Python, AWS, Docker")
        
            # 単金
            project_rate_min = st.number_input("案件単価 (最小)", min_value=0, step=1, placeholder="例: 4500")
            project_rate_max = st.number_input("案件単価 (最大)", min_value=0, step=1, placeholder="例: 5500")
        
            # 最寄り駅
            nearest_station = st.text_input("最寄り駅", placeholder="例: 新宿駅")
        
            # 働き方
            project_remote = st.selectbox("案件働き方", ["リモート可能", "オンサイトのみ"])
        
            # 商流
            business_flow = st.text_input("エンジニア商流 (カンマ区切り)", placeholder="例: 直請け, 一次請け")
            project_business_flow = st.selectbox("案件商流", ["直請け", "一次請け", "二次請け"])
        
            # 年齢制限
            age = st.number_input("エンジニア年齢", min_value=0, step=1, placeholder="例: 30")
            age_min = st.number_input("案件年齢制限 (最小)", min_value=0, step=1, placeholder="例: 25")
            age_max = st.number_input("案件年齢制限 (最大)", min_value=0, step=1, placeholder="例: 35")
        
            # 外国籍
            nationality = st.radio("エンジニア外国籍", ["日本国籍", "外国籍"])
            project_foreign = st.selectbox("案件外国籍対応", ["外国籍不可", "外国籍対応可能"])

            submitted = st.form_submit_button("人材に最適な案件を検索")
        
        if submitted:
            if st.session_state["tab1"]:
                # Google APIサービスの認証
                credentials = authenticate_google_services()
                docs_service = build('docs', 'v1', credentials=credentials)
                    
                # Google DocsとSheetsのデータ取得
                document_text = get_document_text(docs_service, "1HjjtYZ1RCTPSXLxW5ujCviIwCdMi2a4gTB7y5wARht0")
                
                sheet_data = st.session_state["tab1"][:50]
                st.write(f"「{st.session_state[f'search_query_tab1']}」を含む最新 {len(sheet_data)} 件の案件からマッチングしました！")

                # 案件情報をプロンプトに追加
                case_info = "\n".join([
                    f"案件 {i+1}:\n日時: {row[0]}\n件名: {row[3]}\n本文: {row[4]}"
                    for i, row in enumerate(sheet_data)
                ])
    
                # プロンプトの作成
                combined_prompt = (
                    f"プロンプト:\n{document_text}\n\n"
                    f"候補者情報:\n\n"
                    f"**スキル: {skills}\n"
                    f"**最寄り駅: {nearest_station}\n"
                    f"**外国籍: {nationality}\n\n"
                    f"案件条件:\n"
                    f"**単価範囲: {project_rate_min} 円 - {project_rate_max} 円\n"
                    f"**案件働き方: {project_remote}\n"
                    f"**商流: {project_business_flow}\n"
                    f"**年齢制限: {age_min} - {age_max}\n"
                    f"**外国籍対応: {project_foreign}\n"
                    f"案件情報:\n{case_info}"
                )
    
                with st.spinner("AIによるマッチング率計算中..."):
                    # AIにプロンプトを送信
                    response = openai.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": "あなたは案件をマッチングするエージェントです。"},
                            {"role": "user", "content": combined_prompt},
                        ]
                    )
                    raw_response = response.choices[0].message.content.strip()
                
                process_and_render_results(raw_response)
                
            else:
                st.warning("検索結果がありません。")
            
            
    with tab2:
        search_data(
            sheet_id="1amJJDVMr3__OmLgWo1Z9w6FXZ9aMaNm0WRlx1_TYXnE",
            range_name="'【案件】DB'!A:O",
            form_id="tab2"
        )
        # 人材情報入力フォーム
        with st.form("input_form"):
            st.subheader("人材情報入力フォーム")
            email_content = st.text_area("【必須】人材要件のメール文を貼り付け", placeholder="ここにメール文を貼り付けてください")
            skill_sheet = st.file_uploader("【必須】スキルシートアップロード（PDF、Excel、Word形式）", type=["pdf", "xlsx", "xls", "docx"])
            submitted = st.form_submit_button("人材に最適な案件を検索")
        
        if submitted:
            if st.session_state["tab2"]:
                if not skill_sheet or not email_content.strip():
                    st.error("すべての必須項目を入力してください。")
                else:
                    extracted_text = extract_text_from_file(skill_sheet)
                    if extracted_text:
            
                        # Google APIサービスの認証
                        credentials = authenticate_google_services()
                        docs_service = build('docs', 'v1', credentials=credentials)
            
                        # Google DocsとSheetsのデータ取得
                        document_text = get_document_text(docs_service, "1HjjtYZ1RCTPSXLxW5ujCviIwCdMi2a4gTB7y5wARht0")
                        sheet_data = st.session_state["tab2"][:50]
                        st.write(f"「{st.session_state[f'search_query_tab2']}」を含む最新 {len(sheet_data)} 件の案件からマッチングしました！")

                        # 案件情報をプロンプトに追加
                        case_info = "\n".join([
                            f"案件 {i+1}:\n日時: {row[0]}\n件名: {row[3]}\n本文: {row[4]}"
                            for i, row in enumerate(sheet_data)
                        ])
            
                        # プロンプトの作成
                        combined_prompt = (
                            f"プロンプト:\n{document_text}\n\n"
                            f"候補者情報:\n{email_content.strip()}\n\n"
                            f"候補者のスキルシート:\n{extracted_text}\n\n"
                            f"案件情報:\n{case_info}"
                        )
        
                        with st.spinner("AIによるマッチング率計算中..."):
                            # AIにプロンプトを送信
                            response = openai.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[
                                    {"role": "system", "content": "あなたは案件をマッチングするエージェントです。"},
                                    {"role": "user", "content": combined_prompt},
                                ]
                            )
                            raw_response = response.choices[0].message.content.strip()

                        process_and_render_results(raw_response)
                        
    with tab3:
        search_data(
            sheet_id="1amJJDVMr3__OmLgWo1Z9w6FXZ9aMaNm0WRlx1_TYXnE",
            range_name="'【人材】DB'!A:O",
            form_id="tab3"
        )
        # キャッシュデータの初期化
        if st.button("Clear Cache"):
            st.cache_data.clear()
            
        # 情報入力フォーム
        with st.form("input_form_case_detail"):
            st.subheader("情報入力フォーム")
                    
            # エンジニアスキル
            skills = st.text_input("エンジニアスキル (カンマ区切り)", placeholder="例: Python, AWS, Docker")
        
            # 単金
            project_rate_min = st.number_input("案件単価 (最小)", min_value=0, step=1, placeholder="例: 4500")
            project_rate_max = st.number_input("案件単価 (最大)", min_value=0, step=1, placeholder="例: 5500")
        
            # 最寄り駅
            nearest_station = st.text_input("最寄り駅", placeholder="例: 新宿駅")
        
            # 働き方
            project_remote = st.selectbox("案件働き方", ["リモート可能", "オンサイトのみ"])
        
            # 商流
            business_flow = st.text_input("エンジニア商流 (カンマ区切り)", placeholder="例: 直請け, 一次請け")
            project_business_flow = st.selectbox("案件商流", ["直請け", "一次請け", "二次請け"])
        
            # 年齢制限
            age = st.number_input("エンジニア年齢", min_value=0, step=1, placeholder="例: 30")
            age_min = st.number_input("案件年齢制限 (最小)", min_value=0, step=1, placeholder="例: 25")
            age_max = st.number_input("案件年齢制限 (最大)", min_value=0, step=1, placeholder="例: 35")
        
            # 外国籍
            nationality = st.radio("エンジニア外国籍", ["日本国籍", "外国籍"])
            project_foreign = st.selectbox("案件外国籍対応", ["外国籍不可", "外国籍対応可能"])
            submitted = st.form_submit_button("人材に最適な案件を検索")

        if submitted:
            # Google APIサービスの認証
            credentials = authenticate_google_services()
            docs_service = build('docs', 'v1', credentials=credentials)
            drive_service = build('drive', 'v3', credentials=credentials)

            # Google DocsとSheetsのデータ取得
            document_text_true = get_document_text(docs_service, "1XVp9yJgNjmBOKVQZ7KivTMyO899sL9zvQssA_QDZD4I")
            
            # F列がTRUEで直近50件取得
            sheet_data_true = st.session_state["tab3"][:50]
            st.write(f"「{st.session_state[f'search_query_tab3']}」を含む最新 {len(sheet_data_true)} 件の案件からマッチングしました！")
            
            # H列 (URL列) または G列 (ファイルID列) を使用して中身を取得
            combined_data = []
            extracted_files_data = []
            for row in sheet_data_true:
                try:
                    file_id = row[6].strip() if len(row) > 6 else None
            
                    # H列のURLとG列のIDを取得
                    # file_url = row[7] if len(row) > 7 else None
                    g_file_id = row[6] if len(row) > 6 else None
                    
                    # 両方が空の場合はスキップ
                    # if not file_url and not g_file_id:
                    #     continue  # この行をスキップ
            
                    # # H列のURLからファイルIDを抽出
                    # if file_url:
                    #     try:
                    #         file_id = extract_file_id(file_url)  # H列のURLからファイルIDを抽出
                    #     except ValueError:
                    #         pass  # H列のURLが無効の場合は次に進む
            
                    # # H列から抽出できなかった場合、G列のIDを使用
                    # if not file_id and g_file_id:
                    # file_id = g_file_id.strip()  # G列の値をそのまま使用
                    
                    st.write(f"Processing File ID: {file_id}")

                    result = process_file(file_id, drive_service)

                    # 결과 확인
                    if result["file_name"] and result["extracted_text"]:
                        st.write(f"Processed File: {result['file_name']}")
                        st.write(f"Extracted Text: {result['extracted_text'][:500]}")
                    else:
                        st.warning("파일 처리에 실패했습니다.")
                        
                    # request = drive_service.files().get_media(fileId=file_id)
                    # file_content = request.execute()
                    # st.write(f"File content downloaded for ID: {file_id}")
                    
                    # ファイル内容を取得
                    combined_text = f"本文: {row[4]}\nスキルシート内容: {result['extracted_text']}"
                    st.write(combined_text)
                    combined_data.append(combined_text)
                except Exception as e:
                    st.warning(f"ファイル処理中にエラーが発生しました。行データ: {row}, エラー: {e}")
            
           # E列とI列の情報をもとにプロンプトを生成
            case_prompts = [
                f"人材情報 {i+1}:\n{content}" for i, content in enumerate(combined_data)
            ]

            # case_prompts 確認
            st.write("Generated case_prompts:", case_prompts)
            
            # プロンプトに統合
            combined_prompt = (
                f"プロンプト:\n{document_text_true}\n\n"
                f"候補者情報:\n\n"
                f"**スキル: {skills}\n"
                f"**最寄り駅: {nearest_station}\n"
                f"**外国籍: {nationality}\n\n"
                f"案件条件:\n"
                f"**単価範囲: {project_rate_min} 円 - {project_rate_max} 円\n"
                f"**案件働き方: {project_remote}\n"
                f"**商流: {project_business_flow}\n"
                f"**年齢制限: {age_min} - {age_max}\n"
                f"**外国籍対応: {project_foreign}\n"
                f"人材詳細:\n" + "\n".join(case_prompts)
            )
            
            with st.spinner("AIによるマッチング率計算中..."):
                # AIにプロンプトを送信
                response = openai.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "あなたは案件に最適な人材を提案するエージェントです。"},
                        {"role": "user", "content": combined_prompt},
                    ]
                )
                raw_response = response.choices[0].message.content.strip()
                
                process_and_render_results(raw_response)
                    
            st.write(raw_response)
            
    with tab4:
        search_data(
            sheet_id="1amJJDVMr3__OmLgWo1Z9w6FXZ9aMaNm0WRlx1_TYXnE",
            range_name="'【人材】DB'!A:O",
            form_id="tab4"
        )
        # 案件情報入力フォーム
        with st.form("input_form_case"):
            st.subheader("案件情報入力フォーム")
            email_content = st.text_area("【必須】案件概要のメール文を貼り付け", placeholder="ここにメール文を貼り付けてください")
            submitted = st.form_submit_button("案件に最適な人材を検索")
        
        if submitted:
            if not email_content.strip():
                st.error("すべての必須項目を入力してください。")
            else:
                # Google APIサービスの認証
                credentials = authenticate_google_services()
                docs_service = build('docs', 'v1', credentials=credentials)
                drive_service = build('drive', 'v3', credentials=credentials)

                # Google DocsとSheetsのデータ取得
                document_text_true = get_document_text(docs_service, "1XVp9yJgNjmBOKVQZ7KivTMyO899sL9zvQssA_QDZD4I")
                # F列がTRUEで直近50件取得
                sheet_data_true = st.session_state["tab4"][:50]
                st.write(f"「{st.session_state[f'search_query_tab4']}」を含む最新 {len(sheet_data_true)} 件の案件からマッチングしました！")

                # H列 (URL列) または G列 (ファイルID列) を使用して中身を取得
                extracted_files_data = []
                for row in sheet_data_true:
                    try:
                        file_id = None
                
                        # H列のURLとG列のIDを取得
                        file_url = row[7] if len(row) > 7 else None
                        g_file_id = row[6] if len(row) > 6 else None
                
                        # 両方が空の場合はスキップ
                        if not file_url and not g_file_id:
                            continue  # この行をスキップ
                
                        # H列のURLからファイルIDを抽出
                        if file_url:
                            try:
                                file_id = extract_file_id(file_url)  # H列のURLからファイルIDを抽出
                            except ValueError:
                                pass  # H列のURLが無効の場合は次に進む
                
                        # H列から抽出できなかった場合、G列のIDを使用
                        if not file_id and g_file_id:
                            file_id = g_file_id.strip()  # G列の値をそのまま使用
                
                        if file_id:
                            # ファイル内容を取得
                            file_content = get_file_content_from_drive(file_id, drive_service)
                            extracted_files_data.append(file_content)
                        else:
                            extracted_files_data.append("")
                    except Exception as e:
                        st.warning(f"ファイル処理中にエラーが発生しました。行データ: {row}, エラー: {e}")
                        extracted_files_data.append("")
                
                # I列に格納（処理用）
                for i, row in enumerate(sheet_data_true):
                    if i < len(extracted_files_data):
                        row.append(extracted_files_data[i])  # I列にファイルの内容を格納
                
                # E列とI列の情報をもとにプロンプトを生成
                case_prompts = [
                    f"人材情報 {i+1}:\n本文: {row[4]}\nスキルシートファイル内容: {row[8]}"
                    for i, row in enumerate(sheet_data_true)
                ]
                
                # プロンプトに統合
                combined_prompt = (
                    f"プロンプト:\n{document_text_true}\n\n"
                    f"人材情報:\n{email_content.strip()}\n\n"
                    f"人材詳細:\n" + "\n".join(case_prompts)
                )
                
                with st.spinner("AIによるマッチング率計算中..."):
                    # AIにプロンプトを送信
                    response = openai.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": "あなたは案件に最適な人材を提案するエージェントです。"},
                            {"role": "user", "content": combined_prompt},
                        ]
                    )
                    raw_response = response.choices[0].message.content.strip()
                process_and_render_results(raw_response)
